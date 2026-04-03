#!/usr/bin/env python3
"""
Dhwani RIS - Daily RFP Scout Agent
====================================
Scrapes RFPs from DevNetJobsIndia (and Devex), scores them for Dhwani relevance,
drafts full proposals in Dhwani's style, and sends an email digest.

Usage:
    python rfp_agent.py                   # Run once
    python rfp_agent.py --no-email        # Run without sending email
    python rfp_agent.py --test            # Test with 3 RFPs only

Requirements: See requirements.txt
Config: Edit config.yaml before running
"""

import os
import re
import json
import yaml
import time
import smtplib
import logging
import argparse
import requests
from datetime import datetime
from bs4 import BeautifulSoup
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import anthropic
    ANTHROPIC_AVAILABLE = True
except ImportError:
    ANTHROPIC_AVAILABLE = False

# ─────────────────────────────────────────────
#  DHWANI COMPANY KNOWLEDGE BASE
# ─────────────────────────────────────────────

DHWANI_PROFILE = """
COMPANY: Dhwani Rural Information Systems Pvt Ltd (Dhwani RIS)
ISO Certified: ISO 27001:2013
Location: Plot 94, Sector 44, Gurgaon, Haryana 122022
Email: reachus@dhwaniris.com | Website: www.dhwaniris.in
Experience: 10+ years | Team: 130+ ICT4D professionals | Projects: 300+ | Clients: 140+ | Countries: 6

PRODUCTS:
1. mGrant — Grant & Program Lifecycle Management (SaaS, cloud-based)
   - Used by 40+ donors, 1500+ grantees across India
   - Built on Frappe framework (Python-based enterprise framework)
   - Key modules: NGO onboarding/due diligence, proposal & application management,
     grant disbursement tracking, LFA (Logical Framework Approach) tracking,
     M&E and activities audit, CSR compliance (CSR-2 form, Annual Action Plan),
     real-time dashboards & analytics, grantee workspace
   - Integrations: mForm, PowerBI, Tableau, Superset, Google Data Studio

2. mForm — Mobile Data Collection App (offline-enabled)
   - Field-level offline data capture with configurable forms
   - Geo-tagged, time-stamped, media-enabled (photos, signatures, documents)
   - Integrates with mGrant for automated sync
   - Used for program MIS, surveys, beneficiary tracking

3. mLearn — Learning Management System
   - Built on Moodle (open source LMS)
   - Customised for social sector training and capacity building

SERVICES:
- Custom IT Solutions (web & mobile app development for social sector)
- Tech Consulting (digital transformation, MIS design, system architecture)
- DaaS (Dashboard as a Service) — Superset, Power BI, Tableau, Google Data Studio
- Open source implementation: Frappe/ERPNext, Moodle, Zoho, CLIIFE
- M&E system design and deployment
- Health Information Systems (HMIS, patient tracking, nutrition platforms)
- Education Management Information Systems (EMIS, school management)
- Data systems: ODK, KoBoToolbox, CommCare implementation and customisation

KEY CLIENTS:
- CSRs/Corporates: HDFC Bank, Godrej CSR, Axis Bank Foundation, Reliance Foundation,
  Adani Foundation, Amazon, Google, Tata Steel, Mahindra, JSLA, Kyndryl
- Foundations: Tata Trusts, Shiv Nadar Foundation, Azim Premji Foundation, Piramal Foundation
- International: UNICEF, World Bank, WHO, GIZ, USAID-funded programs
- Implementation NGOs: PRADAN, Transform Rural India, AKRSP, Swasti
- Research/Consulting: KPMG, Deloitte, GDi, MSC, EY, CMS

SECTORS: Health, Education, Livelihoods/Agriculture, Gender/Women Empowerment,
         WASH, Environment, CSR/Grant Management, Social Protection, Rural Development

WE ARE THE RIGHT FIT FOR RFPs ASKING FOR:
- Grant management system / software / platform
- Digital platform for NGOs / CSOs / foundations
- MIS (Management Information System) development
- Data collection app / tool / mobile application
- M&E (Monitoring & Evaluation) system / platform
- Health information system / HMIS / patient tracking
- Education management system / EMIS / school platform
- Dashboard and data visualization / analytics
- Technology consulting for development sector
- Custom software development for social impact
- IVRS / IVR-based system development
- CRM for NGOs or development organisations
- Digital transformation consulting for nonprofits
- AI/ML tools for social sector
- ODK / KoBoToolbox / CommCare implementation
- Frappe/ERPNext customisation for social sector
- Program monitoring portal / reporting system
"""

DHWANI_ABOUT_TEXT = """Dhwani Rural Information Systems (ISO Certified), based in Delhi NCR, India, is a distinguished tech advisory and consulting firm with over ten years of experience in the social and development sectors. Our dynamic team of 130+ ICT4D professionals specializes in creating innovative IT applications that drive social impact for NGOs, government agencies, CSRs, and international funding organizations. We pride ourselves on being a diverse group of development professionals, tech enthusiasts, software engineers, and analysts, all united by a common goal: to enable organizations to maximize their social impact. Our cutting-edge solutions empower clients to seamlessly monitor and evaluate their programs, ensuring efficient and effective operations in their mission-driven endeavours."""

DHWANI_CLIENT_TEXT = """Dhwani has worked on more than 400+ projects with clients that range from Funders and Foundations, Grassroots Organisations and CSOs, Collaboratives, Consulting, and Research Organizations, to Governments and multilateral agencies.

We are working with the following funding organizations:
1. CSRs like HDFC Bank, Godrej CSR, Axis Bank Foundation, Reliance Foundation, Hyundai Motors India Foundation, etc.
2. Large-scale funding organizations include the Tata Trusts, UNICEF, World Bank, WHO, and GIZ, among others.
3. Implementation organizations like Transform Rural India, AKRSP, PRADAN, Piramal Foundation, Azim Premji Foundation, etc."""

DHWANI_MGRANT_TEXT = """mGrant is a comprehensive, cloud-based SaaS platform that streamlines the entire grant lifecycle for donor and CSR organizations. It is currently trusted by over 40+ donors and 1,500+ grantee organizations across India. Built on the robust Frappe framework, mGrant ensures scalability, security, and unlimited customisation capabilities.

Key capabilities:
• NGO Onboarding & Due Diligence — Self-registration portal, document uploads, verification, rating & scoring
• Proposal & Application Management — Configurable templates, RFP/TOR release, multi-level assessment, approval workflows
• Comprehensive Grants Management — End-to-end grant lifecycle tracking, budget/fund tracking, LFA-based impact monitoring
• M&E and Activities Audit Module — Evidence submission, geotagged uploads, multi-level approval for verification
• Advanced Dashboards & Analytics — Real-time insights across grants, geographies, programs; multi-dimensional data slicing
• CSR Compliance Management — Compliance tracking for CSR-2 form, Annual Action Plan, automated documentation
• Grantee Workspace — Dedicated portal for grantees to track grants, submit reports, and communicate with donors
• Seamless Integrations — mForm for field data collection; PowerBI, Tableau, Superset for BI; custom API support"""

DHWANI_MFORM_TEXT = """mForm is an offline-enabled, mobile-based data collection application that enables accurate and timely data capture directly from project locations, even in low or no internet connectivity environments.

mForm enables:
• Offline data capture through configurable digital forms
• Geo-tagged and time-stamped data collection for authenticity
• Media capture (photos, documents, signatures) for verification
• Automated data synchronization with mGrant once connectivity is available
• Reduction of manual data entry errors and reporting delays"""

# Keywords to identify relevant vs irrelevant RFPs
RELEVANT_KEYWORDS = [
    "technology", "software", "digital", "platform", "application", "app", "system",
    "data collection", "MIS", "HMIS", "EMIS", "information system", "management system",
    "monitoring", "evaluation", "M&E", "dashboard", "analytics", "reporting system",
    "grant management", "CSR management", "program management", "fund management",
    "mobile", "mhealth", "health tech", "edtech", "education technology", "healthtech",
    "IT solution", "digital solution", "tech platform", "system development", "web application",
    "database", "automation", "digitization", "digitalization", "digital transformation",
    "CRM", "ERP", "portal", "API", "integration", "tool development",
    "capacity building technology", "technology consulting", "IT consulting",
    "IVRS", "IVR", "SMS platform", "WhatsApp", "helpline system", "chatbot",
    "AI", "machine learning", "data science", "GIS", "mapping", "geospatial",
    "ODK", "KoBoToolbox", "CommCare", "DHIS2", "OpenMRS", "Frappe",
    "impact measurement", "beneficiary tracking", "outcome tracking",
    "survey tool", "data management", "cloud", "SaaS", "mobile app",
]

IRRELEVANT_KEYWORDS = [
    "housekeeping", "security guard", "driver", "vehicle hire", "taxi", "cab hiring",
    "travel agency", "hotel booking", "accommodation", "catering", "pantry",
    "construction", "civil work", "plumbing", "electrical", "renovation",
    "printing press", "stationery", "furniture", "equipment procurement", "hardware supply",
    "external audit", "CA firm", "chartered accountant", "statutory audit",
    "documentary film", "photography", "videography", "media production",
    "nukkad natak", "street play", "event management", "logistics",
    "manpower supply", "staffing agency", "outsourced manpower",
    "sewerage", "WASH infrastructure", "water pipeline", "sanitation infrastructure",
    "advocacy", "policy research", "qualitative study", "endline study",
    "baseline assessment", "evaluation study", "impact evaluation",
]


# ─────────────────────────────────────────────
#  SCRAPING
# ─────────────────────────────────────────────

def scrape_devnetjobsindia(test_mode=False):
    """Scrape RFP listings from DevNetJobsIndia"""
    url = "https://www.devnetjobsindia.org/rfp_assignments.aspx"
    headers = {
        'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
    }

    session = requests.Session()
    try:
        r = session.get(url, headers=headers, timeout=25)
        r.raise_for_status()
    except requests.RequestException as e:
        logging.error(f"Failed to fetch DevNetJobsIndia: {e}")
        return []

    soup = BeautifulSoup(r.text, 'html.parser')

    # Extract ASP.NET ViewState for postback
    viewstate_el = soup.find('input', {'id': '__VIEWSTATE'})
    viewstate = viewstate_el['value'] if viewstate_el else ''
    viewstategen_el = soup.find('input', {'id': '__VIEWSTATEGENERATOR'})
    viewstategen = viewstategen_el['value'] if viewstategen_el else ''

    # Find all RFP rows in grid
    all_rows = soup.find_all('tr')
    grid_rows = [row for row in all_rows if row.find('a', href=lambda x: x and 'doPostBack' in str(x))]
    # Remove duplicates (rows that also have logo links)
    grid_rows = [row for row in grid_rows if row.find('span', id=lambda x: x and 'lblJobTitle' in str(x))]

    if test_mode:
        grid_rows = grid_rows[:3]

    rfps = []
    for i, row in enumerate(grid_rows):
        title_el  = row.find('span', id=lambda x: x and 'lblJobTitle' in str(x))
        org_el    = row.find('span', id=lambda x: x and 'lblJobCo' in str(x))
        loc_el    = row.find('span', id=lambda x: x and 'lblLocation' in str(x))
        dead_el   = row.find('span', id=lambda x: x and 'lblDeadline' in str(x))
        sector_el = row.find('span', id=lambda x: x and 'lblSectors' in str(x))

        title    = title_el.get_text(strip=True)  if title_el  else ''
        org      = org_el.get_text(strip=True)    if org_el    else ''
        location = loc_el.get_text(strip=True)    if loc_el    else ''
        deadline = dead_el.get_text(strip=True)   if dead_el   else ''
        sector   = sector_el.get_text(strip=True) if sector_el else ''

        if not title:
            continue

        # Try to get job_id from logo image src (joblogos/XXXXX.png)
        img_el = row.find('img', src=lambda x: x and 'joblogos' in str(x))
        job_id = None
        if img_el:
            match = re.search(r'joblogos/(\d+)', img_el.get('src', ''))
            if match:
                job_id = match.group(1)

        # Postback target to get job detail
        link_el = row.find('a', href=lambda x: x and 'doPostBack' in str(x) and 'lnkJobTitle' in str(x))
        postback_target = ''
        if link_el:
            match = re.search(r"doPostBack\('([^']+)'", link_el.get('href', ''))
            if match:
                postback_target = match.group(1)

        rfps.append({
            'title': title,
            'organization': org,
            'location': location,
            'deadline': deadline,
            'sector': sector,
            'job_id': job_id,
            'postback_target': postback_target,
            'source': 'DevNetJobsIndia',
            'row_index': i,
            'full_description': None,
            'url': None,
        })

    logging.info(f"DevNetJobsIndia: found {len(rfps)} RFP listings")

    # Fetch full descriptions
    post_headers = {**headers,
                    'Content-Type': 'application/x-www-form-urlencoded',
                    'Referer': url}

    for rfp in rfps:
        try:
            if rfp['job_id']:
                detail_url = f"https://devnetjobsindia.org/JobDescription.aspx?Job_Id={rfp['job_id']}"
                resp = session.get(detail_url, headers=headers, timeout=20)
                rfp['url'] = detail_url
                rfp['full_description'] = _extract_devnet_description(resp.text)
            elif rfp['postback_target']:
                post_data = {
                    '__EVENTTARGET':    rfp['postback_target'],
                    '__EVENTARGUMENT': '',
                    '__VIEWSTATE':      viewstate,
                    '__VIEWSTATEGENERATOR': viewstategen,
                }
                resp = session.post(url, data=post_data, headers=post_headers, timeout=20, allow_redirects=True)
                # Extract job_id from redirected URL
                match = re.search(r'Job_Id=(\d+)', resp.url, re.IGNORECASE)
                if match:
                    rfp['job_id'] = match.group(1)
                    rfp['url'] = resp.url
                rfp['full_description'] = _extract_devnet_description(resp.text)
                time.sleep(0.6)  # be polite
        except Exception as e:
            logging.warning(f"Could not fetch detail for '{rfp['title'][:40]}': {e}")

    return rfps


def _extract_devnet_description(html):
    """Extract the main body text from a DevNetJobsIndia job description page"""
    soup = BeautifulSoup(html, 'html.parser')
    # Remove nav, scripts, styles, footer
    for tag in soup.find_all(['nav', 'header', 'footer', 'script', 'style', 'noscript']):
        tag.decompose()

    # The main content is usually inside ContentPlaceHolder1
    main = soup.find('div', id=lambda x: x and 'ContentPlaceHolder1' in str(x))
    if main:
        text = main.get_text(separator='\n', strip=True)
        # Clean up excessive blank lines
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text[:4000]

    return soup.get_text(separator='\n', strip=True)[:3000]


def try_scrape_devex(config):
    """
    Attempt to scrape Devex for RFPs.
    Devex requires a Pro subscription for most content.
    This function scrapes what's publicly visible.
    """
    rfps = []
    # Devex blocks automated requests (403).
    # To enable: add devex_cookie in config.yaml from a logged-in session.
    cookie = config.get('devex_cookie', '')
    if not cookie:
        logging.info("Devex: No session cookie configured. Skipping. (Add devex_cookie to config.yaml)")
        return rfps

    headers = {
        'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36',
        'Cookie': cookie,
    }
    try:
        r = requests.get('https://www.devex.com/jobs/rfps', headers=headers, timeout=20)
        if r.status_code == 200:
            soup = BeautifulSoup(r.text, 'html.parser')
            # Parse Devex job cards
            cards = soup.find_all('div', class_=lambda x: x and 'job' in str(x).lower())
            for card in cards[:20]:
                title_el = card.find(['h2', 'h3', 'a'])
                title = title_el.get_text(strip=True) if title_el else ''
                if title:
                    rfps.append({
                        'title': title,
                        'organization': '',
                        'location': '',
                        'deadline': '',
                        'sector': '',
                        'source': 'Devex',
                        'url': 'https://www.devex.com/jobs/rfps',
                        'full_description': None,
                        'job_id': None,
                    })
    except Exception as e:
        logging.warning(f"Devex scraping failed: {e}")

    logging.info(f"Devex: found {len(rfps)} RFPs")
    return rfps


# ─────────────────────────────────────────────
#  SCORING & FILTERING
# ─────────────────────────────────────────────

def score_rfp(rfp, claude_client):
    """Score RFP relevance for Dhwani. Uses Claude if available, else keyword matching."""
    if claude_client:
        return _score_with_claude(rfp, claude_client)
    return _score_with_keywords(rfp)


def _score_with_claude(rfp, client):
    desc = (rfp.get('full_description') or '')[:2000]
    prompt = f"""You are a business development expert for Dhwani RIS — an Indian IT company serving the social sector.

DHWANI PROFILE:
{DHWANI_PROFILE}

EVALUATE THIS RFP:
Title: {rfp['title']}
Organization: {rfp['organization']}
Sector: {rfp.get('sector', '')}
Location: {rfp.get('location', '')}
Deadline: {rfp.get('deadline', '')}
Description: {desc if desc else 'Not available'}

Score from 0–10 for Dhwani RIS relevance:
10 = Perfect (e.g. grant management software, MIS platform, data collection app, HMIS)
7–9 = Good fit (tech solution in a sector Dhwani serves)
4–6 = Moderate (some tech component but not Dhwani's core)
1–3 = Weak (minimal tech)
0 = Not relevant (physical goods, housekeeping, travel logistics, CA audit, film production, etc.)

Respond ONLY with valid JSON:
{{
  "score": <0-10 integer>,
  "recommendation": "apply" | "consider" | "skip",
  "reason": "<one concise sentence>",
  "relevant_product": "<mGrant | mForm | mLearn | Custom Dev | Tech Consulting | DaaS | None>",
  "key_requirements": ["<req1>", "<req2>", "<req3>"]
}}"""

    try:
        msg = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=400,
            messages=[{"role": "user", "content": prompt}]
        )
        text = msg.content[0].text.strip()
        # Extract JSON from response
        json_match = re.search(r'\{.*\}', text, re.DOTALL)
        if json_match:
            return json.loads(json_match.group())
    except Exception as e:
        logging.warning(f"Claude scoring failed for '{rfp['title'][:40]}': {e}")

    return _score_with_keywords(rfp)


def _score_with_keywords(rfp):
    text = ' '.join([
        rfp.get('title', ''),
        rfp.get('sector', ''),
        rfp.get('organization', ''),
        rfp.get('full_description', '') or ''
    ]).lower()

    score = 0
    for kw in RELEVANT_KEYWORDS:
        if kw.lower() in text:
            score += 1

    for kw in IRRELEVANT_KEYWORDS:
        if kw.lower() in text:
            score -= 3

    score = max(0, min(10, score))
    rec   = 'apply' if score >= 7 else ('consider' if score >= 4 else 'skip')
    return {
        'score': score,
        'recommendation': rec,
        'reason': 'Based on keyword matching (Claude API not configured)',
        'relevant_product': 'Unknown',
        'key_requirements': [],
    }


# ─────────────────────────────────────────────
#  PROPOSAL DRAFTING
# ─────────────────────────────────────────────

def draft_proposal(rfp, scoring, claude_client):
    """Draft a full Dhwani-style proposal for an RFP using Claude."""
    if not claude_client:
        logging.warning("Claude API not configured — proposal drafting skipped.")
        return None

    desc = (rfp.get('full_description') or 'Full description not available')[:3500]
    today = datetime.now().strftime('%d %B %Y')
    relevant_product = scoring.get('relevant_product', 'Custom Development')
    key_requirements = scoring.get('key_requirements', [])

    prompt = f"""You are writing a full proposal on behalf of Dhwani Rural Information Systems Pvt Ltd (Dhwani RIS).

COMPANY BACKGROUND (use this verbatim where appropriate):
ABOUT DHWANI:
{DHWANI_ABOUT_TEXT}

CLIENT PORTFOLIO:
{DHWANI_CLIENT_TEXT}

MGRANT PRODUCT (use when relevant):
{DHWANI_MGRANT_TEXT}

MFORM PRODUCT (use when relevant):
{DHWANI_MFORM_TEXT}

FULL COMPANY PROFILE:
{DHWANI_PROFILE}

────────────────────────────────────────
RFP DETAILS:
Title: {rfp['title']}
Issuing Organization: {rfp['organization']}
Sector: {rfp.get('sector', 'Development Sector')}
Location: {rfp.get('location', 'India')}
Deadline: {rfp.get('deadline', 'As per RFP')}
Source: {rfp.get('url', 'DevNetJobsIndia')}
Key Requirements Identified: {', '.join(key_requirements) if key_requirements else 'See description'}
Most Relevant Dhwani Product/Service: {relevant_product}

FULL RFP DESCRIPTION:
{desc}
────────────────────────────────────────

Write a complete, professional, client-specific proposal following EXACTLY this structure:

# Proposal for [RFP Title]

(Leave cover-page metadata; it will be added as a Word cover page)

## Table of Contents
(List sections with page numbers as placeholders)

## About Dhwani Rural Information Systems
(Use the about text provided above, do not change it)

## Dhwani's Client Portfolio
(Use the client portfolio text above)

## Understanding of the Requirements
(Write 3–5 detailed paragraphs showing deep understanding of the issuing organization's context, challenges, and what they need. Reference specifics from the RFP description. Show empathy and domain knowledge.)

### Desired Outcomes
(List 3–5 bullet outcomes the organization wants to achieve)

## Proposed Solution
(Explain clearly which Dhwani product/service addresses this need and why. If mGrant → explain grant management fit. If mForm → explain data collection fit. If custom dev → describe the custom solution approach. Be specific, not generic. 3–5 paragraphs.)

### Key Platform Features / Solution Capabilities
(List 6–10 specific features/capabilities that address the RFP requirements)

## Implementation Approach
(Describe the phased implementation — use these phases but adapt timing to the context):

### Phase 1: Discovery & Requirements (2–3 weeks)
- Bullet points for this phase

### Phase 2: Configuration & Development (6–10 weeks)
- Bullet points for this phase

### Phase 3: Testing & Training (2–3 weeks)
- Bullet points for this phase

### Phase 4: Go-Live & Support (Ongoing)
- Bullet points for this phase

## Scope of Work
(Detailed deliverables table — list Component | Description as you would in a proposal. 6–10 rows.)

## Budget
[To be finalised based on detailed scoping discussion. Our team will provide a detailed cost breakdown upon further discovery call.]

## Timelines
(High-level Gantt-style timeline in text, mapped to the 4 phases above)

## Warranty & Support
- 90-day warranty post go-live for defect resolution
- Quarterly refresher trainings
- Dedicated helpdesk with ticket-based support
- SLA-based response times (Critical: 4 hours, Major: 8 hours, Minor: 24 hours)
- Annual subscription includes server hosting, maintenance, and future product upgrades

## Governance Mechanism
- Dedicated Project Manager assigned to the engagement
- Weekly status calls during implementation
- Monthly steering committee review
- Escalation matrix with clear ownership at each tier
- Joint project team with [Organization Name] and Dhwani

## Why Dhwani RIS
(Write a compelling 1-page "why us" section highlighting: experience in social sector, relevant product (mGrant/mForm), track record with similar clients, ISO certification, team strength, and support model)

Write in a professional, formal yet warm tone. The proposal should feel tailored specifically to {rfp['organization']}, not generic.
Date of proposal: {today}
IMPORTANT: Do NOT include any pricing/numbers in the Budget section. Leave it as described above."""

    try:
        msg = claude_client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=6000,
            messages=[{"role": "user", "content": prompt}]
        )
        return msg.content[0].text
    except Exception as e:
        logging.error(f"Proposal drafting failed for '{rfp['title'][:40]}': {e}")
        return None


# ─────────────────────────────────────────────
#  WORD DOCUMENT GENERATION
# ─────────────────────────────────────────────

def save_proposal_as_docx(rfp, proposal_text, output_dir):
    """Save the drafted proposal as a Word .docx file"""
    safe_name = re.sub(r'[^\w\s-]', '', rfp['title'])
    safe_name = re.sub(r'\s+', '_', safe_name)[:60]
    filename   = f"{safe_name}_Proposal_Dhwani.docx"
    filepath   = Path(output_dir) / filename

    if not DOCX_AVAILABLE:
        # Fallback: save as plain text
        txt_path = filepath.with_suffix('.txt')
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(f"Proposal: {rfp['title']}\n")
            f.write(f"For: {rfp['organization']}\n")
            f.write(f"Deadline: {rfp.get('deadline','N/A')}\n")
            f.write(f"Source: {rfp.get('url','N/A')}\n\n")
            f.write(proposal_text)
        return str(txt_path)

    doc = Document()

    # ── Cover page ──────────────────────────────
    doc.add_paragraph()  # top spacing

    title_heading = doc.add_heading('', level=0)
    run = title_heading.add_run(f"Proposal for {rfp['title']}")
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    run.font.size = Pt(20)
    title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    for text, bold in [("Submitted To", True), (rfp['organization'], False),
                       ("", False),
                       ("Submitted By", True), ("Dhwani Rural Information Systems Pvt Ltd", False),
                       (datetime.now().strftime('%d %B %Y'), False)]:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        if bold:
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        if text == "Dhwani Rural Information Systems Pvt Ltd":
            run.font.size = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ── Proposal body ────────────────────────────
    for line in proposal_text.split('\n'):
        line_stripped = line.strip()
        if not line_stripped:
            doc.add_paragraph()
            continue

        if line_stripped.startswith('# '):
            doc.add_heading(line_stripped[2:], level=1)
        elif line_stripped.startswith('## '):
            doc.add_heading(line_stripped[3:], level=2)
        elif line_stripped.startswith('### '):
            doc.add_heading(line_stripped[4:], level=3)
        elif line_stripped.startswith('#### '):
            doc.add_heading(line_stripped[5:], level=4)
        elif re.match(r'^[-•*]\s', line_stripped):
            doc.add_paragraph(line_stripped[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s', line_stripped):
            doc.add_paragraph(line_stripped, style='List Number')
        elif line_stripped.startswith('**') and line_stripped.endswith('**'):
            p = doc.add_paragraph()
            p.add_run(line_stripped.strip('**')).bold = True
        else:
            # Handle inline bold (**text**)
            p = doc.add_paragraph()
            parts = re.split(r'\*\*(.+?)\*\*', line_stripped)
            for j, part in enumerate(parts):
                run = p.add_run(part)
                if j % 2 == 1:  # odd indices are bold segments
                    run.bold = True

    # ── Footer ───────────────────────────────────
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.text = (
            "© Dhwani Rural Information Systems Pvt Ltd  |  CONFIDENTIAL  |  "
            "Address: Plot 94, Sector 44, Gurgaon, Haryana 122022  |  "
            "Email: reachus@dhwaniris.com  |  Website: www.dhwaniris.in"
        )
        fp.runs[0].font.size = Pt(7)

    doc.save(str(filepath))
    logging.info(f"  → Saved: {filename}")
    return str(filepath)


# ─────────────────────────────────────────────
#  EMAIL & DIGEST
# ─────────────────────────────────────────────

def send_email_digest(relevant_rfps, config, output_dir):
    """Send the daily RFP digest email with proposal attachments"""
    today_str = datetime.now().strftime('%d %B %Y')
    apply_count   = sum(1 for r in relevant_rfps if r.get('scoring', {}).get('recommendation') == 'apply')
    consider_count = sum(1 for r in relevant_rfps if r.get('scoring', {}).get('recommendation') == 'consider')

    # ── Build HTML body ──────────────────────────
    html = f"""
<!DOCTYPE html>
<html>
<body style="font-family: Arial, sans-serif; max-width: 750px; margin: 0 auto; padding: 20px; color: #333;">

<div style="background: linear-gradient(135deg, #c00000, #8b0000); padding: 20px 25px; border-radius: 8px; margin-bottom: 20px;">
  <h1 style="color: white; margin: 0; font-size: 22px;">🎯 Daily RFP Scout</h1>
  <p style="color: #ffcccc; margin: 6px 0 0 0; font-size: 14px;">Dhwani RIS Business Development &nbsp;|&nbsp; {today_str}</p>
</div>

<div style="background: #f7f7f7; border-left: 4px solid #c00000; padding: 12px 16px; margin-bottom: 20px; border-radius: 0 6px 6px 0;">
  <strong>Today's Summary:</strong>
  &nbsp; ✅ <strong>{apply_count}</strong> to Apply
  &nbsp; 🤔 <strong>{consider_count}</strong> to Consider
  &nbsp; 📄 <strong>{sum(1 for r in relevant_rfps if r.get('proposal_file'))}</strong> Proposals Drafted
  <br><small style="color: #666;">Sources: DevNetJobsIndia &nbsp;|&nbsp; Devex</small>
</div>
"""

    for rfp in relevant_rfps:
        sc   = rfp.get('scoring', {})
        score = sc.get('score', 0)
        rec   = sc.get('recommendation', 'consider')
        reason = sc.get('reason', '')
        product = sc.get('relevant_product', '')
        has_proposal = bool(rfp.get('proposal_file'))

        color  = '#1a7a1a' if score >= 8 else ('#e07800' if score >= 5 else '#c00000')
        badge  = {'apply': '✅ APPLY NOW', 'consider': '🤔 CONSIDER', 'skip': '⏭ SKIP'}.get(rec, rec.upper())

        html += f"""
<div style="border: 1px solid #e0e0e0; border-radius: 8px; padding: 16px; margin-bottom: 16px; box-shadow: 0 1px 3px rgba(0,0,0,0.07);">
  <div style="display: flex; justify-content: space-between; align-items: flex-start; flex-wrap: wrap; gap: 8px;">
    <h3 style="margin: 0; font-size: 16px; color: #1a1a1a; flex: 1;">{rfp['title']}</h3>
    <span style="background:{color}; color:white; padding:4px 12px; border-radius:20px; font-size:13px; font-weight:bold; white-space:nowrap;">{score}/10</span>
  </div>
  <table style="width:100%; margin-top:10px; font-size:13px; border-collapse:collapse;">
    <tr><td style="color:#666; width:130px; padding:3px 0;">Organization</td><td><strong>{rfp['organization']}</strong></td></tr>
    <tr><td style="color:#666; padding:3px 0;">Location</td><td>{rfp.get('location','N/A')}</td></tr>
    <tr><td style="color:#666; padding:3px 0;">Deadline</td><td><strong style="color:#c00000;">{rfp.get('deadline','N/A')}</strong></td></tr>
    <tr><td style="color:#666; padding:3px 0;">Source</td><td>{rfp.get('source','N/A')} &nbsp; <a href="{rfp.get('url','#')}" style="color:#0066cc;">View RFP →</a></td></tr>
    <tr><td style="color:#666; padding:3px 0;">Relevant Product</td><td>{product}</td></tr>
  </table>
  <div style="margin-top:10px; padding:8px 12px; background:#f9f9f9; border-radius:5px; font-size:13px;">
    <strong>{badge}</strong> — {reason}
  </div>
  {'<div style="margin-top:8px; color:#1a7a1a; font-size:13px;">📄 Proposal draft attached as Word document</div>' if has_proposal else ''}
</div>
"""

    html += f"""
<div style="margin-top:24px; padding:14px; background:#f0f0f0; border-radius:6px; font-size:11px; color:#888; text-align:center;">
  Automated daily digest by <strong>Dhwani RIS RFP Scout</strong><br>
  Proposal drafts are attached as Word documents where generated.<br>
  Contact: reachus@dhwaniris.com | www.dhwaniris.in
</div>
</body></html>"""

    # Save HTML digest locally
    digest_path = Path(output_dir) / f"rfp_digest_{datetime.now().strftime('%Y-%m-%d')}.html"
    with open(digest_path, 'w', encoding='utf-8') as f:
        f.write(html)
    logging.info(f"HTML digest saved: {digest_path}")

    # ── Send email ───────────────────────────────
    # Environment variables override config.yaml (used by GitHub Actions)
    email_cfg = config.get('email', {})
    env_sender     = os.environ.get('EMAIL_SENDER')
    env_password   = os.environ.get('EMAIL_PASSWORD')
    env_recipients = os.environ.get('EMAIL_RECIPIENTS')  # comma-separated
    env_smtp_host  = os.environ.get('EMAIL_SMTP_HOST')
    env_smtp_port  = os.environ.get('EMAIL_SMTP_PORT')

    if env_sender:    email_cfg['sender_email'] = env_sender
    if env_password:  email_cfg['password']     = env_password
    if env_recipients:
        email_cfg['recipients'] = [r.strip() for r in env_recipients.split(',')]
    if env_smtp_host: email_cfg['smtp_host']    = env_smtp_host
    if env_smtp_port: email_cfg['smtp_port']    = int(env_smtp_port)
    # Auto-enable if env vars are present
    if env_sender and env_password and env_recipients:
        email_cfg['enabled'] = True

    if not email_cfg.get('enabled'):
        logging.info("Email not enabled. Digest saved to output folder only.")
        return str(digest_path)

    msg = MIMEMultipart('related')
    msg['Subject'] = f"[RFP Scout] {apply_count} to Apply, {consider_count} to Consider — {today_str}"
    msg['From']    = email_cfg['sender_email']
    msg['To']      = ', '.join(email_cfg['recipients'])

    msg.attach(MIMEText(html, 'html'))

    # Attach proposal Word docs
    for rfp in relevant_rfps:
        pf = rfp.get('proposal_file')
        if pf and Path(pf).exists():
            with open(pf, 'rb') as f:
                part = MIMEApplication(
                    f.read(),
                    _subtype='vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
                part.add_header('Content-Disposition', 'attachment',
                                 filename=Path(pf).name)
                msg.attach(part)

    try:
        smtp_host = email_cfg.get('smtp_host', 'smtp.gmail.com')
        smtp_port = email_cfg.get('smtp_port', 587)
        with smtplib.SMTP(smtp_host, smtp_port) as server:
            server.ehlo()
            server.starttls()
            server.login(email_cfg['sender_email'], email_cfg['password'])
            server.sendmail(email_cfg['sender_email'],
                            email_cfg['recipients'],
                            msg.as_string())
        logging.info(f"✉ Email sent to: {email_cfg['recipients']}")
    except Exception as e:
        logging.error(f"Email send failed: {e}")
        logging.info(f"Digest still saved locally at: {digest_path}")

    return str(digest_path)


# ─────────────────────────────────────────────
#  MAIN PIPELINE
# ─────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description='Dhwani RIS Daily RFP Scout')
    parser.add_argument('--no-email',    action='store_true', help='Skip sending email')
    parser.add_argument('--test',        action='store_true', help='Test mode: process first 3 RFPs only')
    parser.add_argument('--no-proposals',action='store_true', help='Skip proposal drafting')
    parser.add_argument('--min-score',   type=int, default=5, help='Min score to include (default: 5)')
    args = parser.parse_args()

    # ── Logging ──────────────────────────────────
    log_fmt = '%(asctime)s [%(levelname)s] %(message)s'
    logging.basicConfig(
        level=logging.INFO,
        format=log_fmt,
        handlers=[
            logging.FileHandler('rfp_agent.log', encoding='utf-8'),
            logging.StreamHandler()
        ]
    )

    # ── Load config ──────────────────────────────
    config_path = Path(__file__).parent / 'config.yaml'
    config = {}
    if config_path.exists():
        with open(config_path, 'r') as f:
            config = yaml.safe_load(f) or {}
    else:
        logging.warning("config.yaml not found — using defaults. Copy config.example.yaml to config.yaml.")

    if args.no_email:
        config.setdefault('email', {})['enabled'] = False

    # ── Output directory ─────────────────────────
    base_out = Path(config.get('output_dir', './rfp_output'))
    today_out = base_out / datetime.now().strftime('%Y-%m-%d')
    today_out.mkdir(parents=True, exist_ok=True)

    # ── Claude client ─────────────────────────────
    api_key = config.get('anthropic_api_key') or os.environ.get('ANTHROPIC_API_KEY')
    claude_client = None
    if ANTHROPIC_AVAILABLE and api_key:
        claude_client = anthropic.Anthropic(api_key=api_key)
        logging.info("Claude API: connected ✓")
    else:
        logging.warning("Claude API not configured — keyword scoring only, no proposal drafting.")

    # ── Banner ────────────────────────────────────
    logging.info("=" * 65)
    logging.info(f"  Dhwani RIS RFP Scout  |  {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    if args.test:
        logging.info("  MODE: TEST (first 3 RFPs only)")
    logging.info("=" * 65)

    # ── Scrape ───────────────────────────────────
    all_rfps = []

    logging.info("▶ Scraping DevNetJobsIndia...")
    devnet = scrape_devnetjobsindia(test_mode=args.test)
    all_rfps.extend(devnet)
    logging.info(f"  → {len(devnet)} RFPs fetched")

    logging.info("▶ Trying Devex...")
    devex = try_scrape_devex(config.get('devex', {}))
    all_rfps.extend(devex)

    logging.info(f"Total RFPs to evaluate: {len(all_rfps)}")

    # ── Score ────────────────────────────────────
    logging.info("▶ Scoring RFPs for Dhwani relevance...")
    relevant = []

    for rfp in all_rfps:
        logging.info(f"  Evaluating: {rfp['title'][:60]}")
        scoring = score_rfp(rfp, claude_client)
        rfp['scoring'] = scoring
        score = scoring.get('score', 0)
        rec   = scoring.get('recommendation', 'skip')
        logging.info(f"    Score: {score}/10  |  {rec.upper()}  |  {scoring.get('reason','')}")

        min_score = args.min_score or config.get('min_score', 5)
        if rec in ('apply', 'consider') and score >= min_score:
            relevant.append(rfp)

        time.sleep(0.3)

    logging.info(f"\n✓ {len(relevant)} relevant RFP(s) found (score ≥ {args.min_score or config.get('min_score', 5)})")

    # ── Draft proposals ───────────────────────────
    draft_threshold = config.get('draft_threshold', 7)

    if not args.no_proposals and claude_client:
        logging.info(f"▶ Drafting proposals for RFPs with score ≥ {draft_threshold}...")
        for rfp in relevant:
            if rfp.get('scoring', {}).get('score', 0) >= draft_threshold:
                logging.info(f"  Drafting: {rfp['title'][:55]}...")
                proposal_text = draft_proposal(rfp, rfp['scoring'], claude_client)
                if proposal_text:
                    filepath = save_proposal_as_docx(rfp, proposal_text, str(today_out))
                    rfp['proposal_file'] = filepath
                time.sleep(1.5)

    # ── Email digest ──────────────────────────────
    if relevant:
        logging.info("▶ Sending email digest...")
        digest_path = send_email_digest(relevant, config, str(today_out))
        logging.info(f"  Digest: {digest_path}")
    else:
        logging.info("No relevant RFPs today — no digest sent.")

    # ── Save JSON log ─────────────────────────────
    log_data = {
        'run_date': datetime.now().isoformat(),
        'total_scraped': len(all_rfps),
        'relevant_count': len(relevant),
        'proposals_drafted': sum(1 for r in relevant if r.get('proposal_file')),
        'rfps': [{
            'title':        r['title'],
            'organization': r['organization'],
            'deadline':     r.get('deadline'),
            'score':        r.get('scoring', {}).get('score'),
            'recommendation': r.get('scoring', {}).get('recommendation'),
            'reason':       r.get('scoring', {}).get('reason'),
            'url':          r.get('url'),
            'has_proposal': bool(r.get('proposal_file')),
        } for r in all_rfps]
    }
    log_file = today_out / 'rfp_run_log.json'
    with open(log_file, 'w', encoding='utf-8') as f:
        json.dump(log_data, f, indent=2, ensure_ascii=False)

    # ── Summary ───────────────────────────────────
    logging.info("\n" + "=" * 65)
    logging.info("  DONE")
    logging.info(f"  Total scraped  : {len(all_rfps)}")
    logging.info(f"  Relevant       : {len(relevant)}")
    logging.info(f"  Proposals made : {sum(1 for r in relevant if r.get('proposal_file'))}")
    logging.info(f"  Output folder  : {today_out}")
    logging.info("=" * 65)

    return relevant


if __name__ == '__main__':
    main()
