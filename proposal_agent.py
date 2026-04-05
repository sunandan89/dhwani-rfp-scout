#!/usr/bin/env python3
"""
proposal_agent.py  —  Dhwani RIS Proposal Drafting Agent
=========================================================
Responsibility: Read scored RFP results, draft full proposals for
top-scoring RFPs using Claude Sonnet, and save as branded Word documents.

This agent is intentionally decoupled from the scout. It can be run:
  • Automatically after rfp_scout.py (pipeline mode)
  • Manually for a specific RFP (interactive mode)
  • On any past results JSON

Usage:
    # Draft proposals for today's scouted RFPs (most common)
    python proposal_agent.py --from-results rfp_output/2026-04-04/rfp_results.json

    # Lower the threshold to draft more proposals
    python proposal_agent.py --from-results rfp_output/.../rfp_results.json --min-score 6

    # Draft a proposal for a single specific RFP by title keyword
    python proposal_agent.py --from-results rfp_output/.../rfp_results.json --only "AgriVaani"

    # Email the drafted proposals after generating them
    python proposal_agent.py --from-results rfp_output/.../rfp_results.json --email

Output:
    rfp_output/YYYY-MM-DD/proposals/<title>_Proposal_Dhwani.docx

Requirements:
    pip install requests pyyaml anthropic python-docx
"""

import os
import re
import json
import yaml
import time
import smtplib
import logging
import argparse
from datetime import datetime
from pathlib import Path
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not installed — proposals will be saved as .txt")

try:
    import anthropic
    ANTHROPIC_AVAILABLE = True
except ImportError:
    ANTHROPIC_AVAILABLE = False

from dhwani_profile import (
    DHWANI_PROFILE,
    DHWANI_ABOUT_TEXT,
    DHWANI_CLIENT_TEXT,
    DHWANI_MGRANT_TEXT,
    DHWANI_MFORM_TEXT,
)


# ─────────────────────────────────────────────
#  PROPOSAL DRAFTING (Claude Sonnet)
# ─────────────────────────────────────────────

def draft_proposal(rfp, scoring, claude_client):
    """
    Draft a full Dhwani-style proposal for a given RFP using Claude Sonnet.
    Returns the proposal as a markdown-formatted string, or None on failure.
    """
    if not claude_client:
        logging.warning("Claude API not configured — cannot draft proposals.")
        return None

    desc = (rfp.get('full_description') or 'Full description not available')[:3500]
    today = datetime.now().strftime('%d %B %Y')
    relevant_product  = scoring.get('relevant_product', 'Custom Development')
    key_requirements  = scoring.get('key_requirements', [])

    prompt = f"""You are writing a full proposal on behalf of Dhwani Rural Information Systems Pvt Ltd (Dhwani RIS).

COMPANY BACKGROUND (use this verbatim where appropriate):

ABOUT DHWANI:
{DHWANI_ABOUT_TEXT}

CLIENT PORTFOLIO:
{DHWANI_CLIENT_TEXT}

MGRANT PRODUCT (use when relevant):
{DHWANI_MGRANT_TEXT}

MFORM PRODUCT (use when relevant):
{DHWANI_MFORM_TEXT}

FULL COMPANY PROFILE:
{DHWANI_PROFILE}

────────────────────────────────────────
RFP DETAILS:
Title: {rfp['title']}
Issuing Organization: {rfp['organization']}
Sector: {rfp.get('sector', 'Development Sector')}
Location: {rfp.get('location', 'India')}
Deadline: {rfp.get('deadline', 'As per RFP')}
Source: {rfp.get('url', 'DevNetJobsIndia')}
Key Requirements Identified: {', '.join(key_requirements) if key_requirements else 'See description'}
Most Relevant Dhwani Product/Service: {relevant_product}

FULL RFP DESCRIPTION:
{desc}
────────────────────────────────────────

Write a complete, professional, client-specific proposal following EXACTLY this structure:

# Proposal for [RFP Title]

## Table of Contents
(List sections with page numbers as placeholders)

## About Dhwani Rural Information Systems
(Use the about text provided above verbatim)

## Dhwani's Client Portfolio
(Use the client portfolio text above)

## Understanding of the Requirements
(3–5 detailed paragraphs showing deep understanding of the organization's context and needs.
Reference specifics from the RFP. Show empathy and domain knowledge.)

### Desired Outcomes
(3–5 bullet outcomes the organization wants to achieve)

## Proposed Solution
(Explain which Dhwani product/service addresses this need. Be specific. 3–5 paragraphs.)

### Key Platform Features / Solution Capabilities
(6–10 specific features/capabilities addressing the RFP requirements)

## Implementation Approach

### Phase 1: Discovery & Requirements (2–3 weeks)
- Bullet points

### Phase 2: Configuration & Development (6–10 weeks)
- Bullet points

### Phase 3: Testing & Training (2–3 weeks)
- Bullet points

### Phase 4: Go-Live & Support (Ongoing)
- Bullet points

## Scope of Work
(Deliverables table — Component | Description, 6–10 rows)

## Budget
[To be finalised based on detailed scoping discussion. Our team will provide a detailed cost breakdown upon further discovery call.]

## Timelines
(High-level timeline mapped to the 4 phases above)

## Warranty & Support
- 90-day warranty post go-live for defect resolution
- Quarterly refresher trainings
- Dedicated helpdesk with ticket-based support
- SLA-based response times (Critical: 4 hours, Major: 8 hours, Minor: 24 hours)
- Annual subscription includes server hosting, maintenance, and future product upgrades

## Governance Mechanism
- Dedicated Project Manager assigned to the engagement
- Weekly status calls during implementation
- Monthly steering committee review
- Escalation matrix with clear ownership at each tier
- Joint project team with [Organization Name] and Dhwani

## Why Dhwani RIS
(Compelling "why us" section: experience in social sector, relevant product track record,
ISO certification, team strength, support model — 1 page)

Write in a professional, formal yet warm tone. Tailor specifically to {rfp['organization']}.
Date of proposal: {today}
IMPORTANT: Do NOT include any actual numbers in the Budget section."""

    try:
        msg = claude_client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=6000,
            messages=[{"role": "user", "content": prompt}]
        )
        return msg.content[0].text
    except Exception as e:
        logging.error(f"Proposal drafting failed for '{rfp['title'][:40]}': {e}")
        return None


# ─────────────────────────────────────────────
#  WORD DOCUMENT GENERATION
# ─────────────────────────────────────────────

def save_proposal_as_docx(rfp, proposal_text, output_dir):
    """Save the drafted proposal as a branded Word .docx file."""
    safe_name = re.sub(r'[^\w\s-]', '', rfp['title'])
    safe_name = re.sub(r'\s+', '_', safe_name)[:60]
    filename  = f"{safe_name}_Proposal_Dhwani.docx"
    filepath  = Path(output_dir) / filename

    if not DOCX_AVAILABLE:
        txt_path = filepath.with_suffix('.txt')
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(f"Proposal: {rfp['title']}\n")
            f.write(f"For: {rfp['organization']}\n")
            f.write(f"Deadline: {rfp.get('deadline','N/A')}\n\n")
            f.write(proposal_text)
        logging.info(f"  → Saved (txt fallback): {txt_path.name}")
        return str(txt_path)

    doc = Document()

    # ── Cover page ──────────────────────────────────────────────
    doc.add_paragraph()
    title_heading = doc.add_heading('', level=0)
    run = title_heading.add_run(f"Proposal for {rfp['title']}")
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    run.font.size = Pt(20)
    title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    cover_lines = [
        ("Submitted To", True),
        (rfp['organization'], False),
        ("", False),
        ("Submitted By", True),
        ("Dhwani Rural Information Systems Pvt Ltd", False),
        (datetime.now().strftime('%d %B %Y'), False),
    ]
    for text, bold in cover_lines:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        if bold:
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        if text == "Dhwani Rural Information Systems Pvt Ltd":
            run.font.size = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ── Proposal body ────────────────────────────────────────────
    for line in proposal_text.split('\n'):
        line_stripped = line.strip()
        if not line_stripped:
            doc.add_paragraph()
            continue
        if line_stripped.startswith('#### '):
            doc.add_heading(line_stripped[5:], level=4)
        elif line_stripped.startswith('### '):
            doc.add_heading(line_stripped[4:], level=3)
        elif line_stripped.startswith('## '):
            doc.add_heading(line_stripped[3:], level=2)
        elif line_stripped.startswith('# '):
            doc.add_heading(line_stripped[2:], level=1)
        elif re.match(r'^[-•*]\s', line_stripped):
            doc.add_paragraph(line_stripped[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s', line_stripped):
            doc.add_paragraph(line_stripped, style='List Number')
        else:
            p = doc.add_paragraph()
            parts = re.split(r'\*\*(.+?)\*\*', line_stripped)
            for j, part in enumerate(parts):
                run = p.add_run(part)
                if j % 2 == 1:
                    run.bold = True

    # ── Footer ───────────────────────────────────────────────────
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.text = (
            "© Dhwani Rural Information Systems Pvt Ltd  |  CONFIDENTIAL  |  "
            "Plot 94, Sector 44, Gurgaon, Haryana 122022  |  "
            "reachus@dhwaniris.com  |  www.dhwaniris.in"
        )
        fp.runs[0].font.size = Pt(7)

    doc.save(str(filepath))
    logging.info(f"  → Saved: {filename}")
    return str(filepath)


# ─────────────────────────────────────────────
#  EMAIL PROPOSALS
# ─────────────────────────────────────────────

def email_proposals(drafted_rfps, config):
    """Send an email with the drafted proposal Word documents attached."""
    today_str = datetime.now().strftime('%d %B %Y')
    count = len(drafted_rfps)

    html = f"""<!DOCTYPE html>
<html><body style="font-family: Arial, sans-serif; max-width: 700px; margin: 0 auto; padding: 20px;">
<div style="background: linear-gradient(135deg, #c00000, #8b0000); padding: 20px 25px; border-radius: 8px; margin-bottom: 20px;">
  <h1 style="color:white; margin:0; font-size:22px;">📄 Proposal Drafts Ready</h1>
  <p style="color:#ffcccc; margin:6px 0 0 0; font-size:14px;">Dhwani RIS  |  {today_str}</p>
</div>
<p><strong>{count} proposal draft(s)</strong> have been generated and are attached to this email.</p>
<ul>
{"".join(f"<li><strong>{r['title']}</strong> — {r['organization']} (Score: {r.get('scoring',{{}}).get('score','?')}/10)</li>" for r in drafted_rfps)}
</ul>
<p style="color:#666; font-size:13px;">These are AI-drafted proposals. Please review, customise, and add specific pricing before sending to clients.</p>
<hr>
<p style="font-size:11px; color:#999;">Generated by Dhwani RIS Proposal Agent | reachus@dhwaniris.com | www.dhwaniris.in</p>
</body></html>"""

    email_cfg = config.get('email', {})
    env_sender     = os.environ.get('EMAIL_SENDER')
    env_password   = os.environ.get('EMAIL_PASSWORD')
    env_recipients = os.environ.get('EMAIL_RECIPIENTS')
    env_smtp_host  = os.environ.get('EMAIL_SMTP_HOST')
    env_smtp_port  = os.environ.get('EMAIL_SMTP_PORT')

    if env_sender:     email_cfg['sender_email'] = env_sender
    if env_password:   email_cfg['password']     = env_password
    if env_recipients: email_cfg['recipients']   = [r.strip() for r in env_recipients.split(',')]
    if env_smtp_host:  email_cfg['smtp_host']    = env_smtp_host
    if env_smtp_port:  email_cfg['smtp_port']    = int(env_smtp_port)
    if env_sender and env_password and env_recipients:
        email_cfg['enabled'] = True

    if not email_cfg.get('enabled'):
        logging.info("Email not enabled — proposals saved to output folder only.")
        return

    msg = MIMEMultipart()
    msg['Subject'] = f"[Proposal Agent] {count} proposal draft(s) ready — {today_str}"
    msg['From']    = email_cfg['sender_email']
    msg['To']      = ', '.join(email_cfg['recipients'])
    msg.attach(MIMEText(html, 'html'))

    for rfp in drafted_rfps:
        pf = rfp.get('proposal_file')
        if pf and Path(pf).exists():
            with open(pf, 'rb') as f:
                part = MIMEApplication(
                    f.read(),
                    _subtype='vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
                part.add_header('Content-Disposition', 'attachment', filename=Path(pf).name)
                msg.attach(part)

    try:
        smtp_host = email_cfg.get('smtp_host', 'smtp.office365.com')
        smtp_port = email_cfg.get('smtp_port', 587)
        with smtplib.SMTP(smtp_host, smtp_port) as server:
            server.ehlo()
            server.starttls()
            server.login(email_cfg['sender_email'], email_cfg['password'])
            server.sendmail(email_cfg['sender_email'], email_cfg['recipients'], msg.as_string())
        logging.info(f"✉ Proposals emailed to: {email_cfg['recipients']}")
    except Exception as e:
        logging.error(f"Email send failed: {e}")


# ─────────────────────────────────────────────
#  MAIN
# ─────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description='Dhwani RIS Proposal Drafting Agent')
    parser.add_argument('--from-results', required=True,
                        help='Path to rfp_results.json from rfp_scout.py')
    parser.add_argument('--min-score', type=int, default=None,
                        help='Minimum score to draft a proposal for (default from config: 7)')
    parser.add_argument('--only', type=str, default=None,
                        help='Only draft proposal for RFPs whose title contains this keyword')
    parser.add_argument('--email', action='store_true',
                        help='Email the drafted proposals when done')
    args = parser.parse_args()

    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s [%(levelname)s] %(message)s',
        handlers=[
            logging.FileHandler('proposal_agent.log', encoding='utf-8'),
            logging.StreamHandler()
        ]
    )

    # Load config
    config_path = Path(__file__).parent / 'config.yaml'
    config = {}
    if config_path.exists():
        with open(config_path) as f:
            config = yaml.safe_load(f) or {}

    # Claude client (required for proposal drafting)
    api_key = config.get('anthropic_api_key') or os.environ.get('ANTHROPIC_API_KEY')
    claude_client = None
    if ANTHROPIC_AVAILABLE and api_key and 'YOUR_ANTHROPIC' not in api_key:
        claude_client = anthropic.Anthropic(api_key=api_key)
        logging.info("Claude API: connected ✓")
    else:
        logging.error(
            "Claude API key required for proposal drafting.\n"
            "Set ANTHROPIC_API_KEY environment variable or add to config.yaml."
        )
        return

    # Load RFP results from scout
    results_path = Path(args.from_results)
    if not results_path.exists():
        logging.error(f"Results file not found: {results_path}")
        return

    with open(results_path, encoding='utf-8') as f:
        all_rfps = json.load(f)

    logging.info(f"Loaded {len(all_rfps)} RFPs from {results_path}")

    # Determine output directory (same date folder as results)
    output_dir = results_path.parent / 'proposals'
    output_dir.mkdir(parents=True, exist_ok=True)

    # Determine draft threshold
    draft_threshold = args.min_score if args.min_score is not None else config.get('draft_threshold', 7)

    # Filter RFPs to draft
    to_draft = [
        r for r in all_rfps
        if r.get('scoring', {}).get('score', 0) >= draft_threshold
        and r.get('scoring', {}).get('recommendation') in ('apply', 'consider')
    ]

    if args.only:
        keyword = args.only.lower()
        to_draft = [r for r in to_draft if keyword in r['title'].lower()]
        logging.info(f"Filtered to RFPs matching '{args.only}': {len(to_draft)}")

    logging.info("=" * 65)
    logging.info(f"  Dhwani RIS Proposal Agent  |  {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    logging.info(f"  Draft threshold: score ≥ {draft_threshold}")
    logging.info(f"  RFPs to draft: {len(to_draft)}")
    logging.info("=" * 65)

    if not to_draft:
        logging.info("No RFPs meet the draft threshold. Done.")
        return

    # ── Draft proposals ─────────────────────────────────────
    drafted = []
    for rfp in to_draft:
        scoring = rfp.get('scoring', {})
        title   = rfp['title'][:60]
        score   = scoring.get('score', 0)
        logging.info(f"▶ Drafting: {title}  (Score: {score}/10)")

        proposal_text = draft_proposal(rfp, scoring, claude_client)
        if proposal_text:
            filepath = save_proposal_as_docx(rfp, proposal_text, str(output_dir))
            rfp['proposal_file'] = filepath
            drafted.append(rfp)
        else:
            logging.warning(f"  ✗ Draft failed for: {title}")

        time.sleep(1.5)  # respect API rate limits

    # ── Email proposals (optional) ──────────────────────────
    if args.email and drafted:
        logging.info("▶ Emailing proposals...")
        email_proposals(drafted, config)

    # ── Summary ─────────────────────────────────────────────
    logging.info("\n" + "=" * 65)
    logging.info("  PROPOSAL AGENT COMPLETE")
    logging.info(f"  Proposals drafted : {len(drafted)} / {len(to_draft)}")
    logging.info(f"  Output folder     : {output_dir}")
    for r in drafted:
        logging.info(f"    • {Path(r['proposal_file']).name}")
    logging.info("=" * 65)


if __name__ == '__main__':
    main()
